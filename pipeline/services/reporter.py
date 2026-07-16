import re
from io import BytesIO

from django.db.models import Count, FloatField, Value
from django.db.models.functions import Coalesce
from django.utils import timezone

from pipeline.models import FeedbackRecord, ProcessingJob


CONSEQUENCE_LABELS = ("Correction", "Improvement", "Prioritization")


def build_executive_report_docx(job: ProcessingJob, feedbacks) -> BytesIO:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"Biblioteca python-docx indisponivel: {exc}") from exc

    document = Document()
    _configure_document(document)

    run_date = timezone.localtime(job.finished_at or timezone.now()).strftime("%d/%m/%Y %H:%M")
    total_feedbacks = feedbacks.count()
    consequence_counts = _consequence_counts(feedbacks)
    negative_count = feedbacks.filter(sentiment_score__lt=-0.05).count()
    negative_percent = _percent(negative_count, total_feedbacks)

    document.add_heading("Relatorio Executivo FEED-ON", level=1)
    document.add_paragraph(f"Data do processamento: {run_date}")
    document.add_paragraph(f"Lote analisado: Job #{job.id} - {job.original_filename}")
    document.add_paragraph(f"Dominio do software: {job.domain_name or 'geral'}")
    document.add_paragraph("As classificacoes sao resultados automaticos e devem ser validadas por uma pessoa responsavel.")
    document.add_paragraph(f"Reasoner Pellet: {'sucesso' if (job.metadata or {}).get('reasoner', {}).get('success') else 'nao executado ou sem sucesso'}")

    document.add_heading("Indicadores Globais", level=2)
    _add_indicators_table(document, total_feedbacks, negative_percent, consequence_counts)

    document.add_heading("Distribuicao de Consequencias", level=2)
    _add_consequence_distribution_table(document, consequence_counts, total_feedbacks)

    top_targets = _top_critical_targets(feedbacks)
    document.add_heading("Top Features Criticas", level=2)
    if top_targets:
        _add_simple_table(document, ("Alvo Semantico", "Corrections"), top_targets)
    else:
        document.add_paragraph("Nao ha features criticas para os filtros selecionados.")

    sentiment_rows = _sentiment_by_target(feedbacks)
    document.add_heading("Media de Sentimento por Categoria", level=2)
    if sentiment_rows:
        _add_simple_table(document, ("Alvo Semantico", "Sentimento Medio"), sentiment_rows)
    else:
        document.add_paragraph("Nao ha dados de sentimento para os filtros selecionados.")

    document.add_heading("Top 10 Critical Issues", level=2)
    critical_rows = list(
        feedbacks.filter(consequence="Correction")
        .annotate(sentiment_sort=Coalesce("sentiment_score", Value(1.0), output_field=FloatField()))
        .order_by("sentiment_sort", "id")[:10]
    )
    if critical_rows:
        _add_critical_issues_table(document, critical_rows)
    else:
        document.add_paragraph("Nao ha issues criticos para os filtros selecionados.")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
        if paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def sanitize_feedback_text(value: str) -> str:
    text = str(value or "")
    text = text.replace('""', '"')
    text = text.replace("\r", " ").replace("\n", " ").replace("\t", " ")
    text = re.sub(r"\s+", " ", text)
    text = text.strip(" []\"';")
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = text.strip()
    if text.endswith("["):
        text = text[:-1].rstrip()
    if text and text[-1] not in ".!?":
        text = f"{text}."
    return text


def _configure_document(document) -> None:
    section = document.sections[0]
    section.top_margin = _inches(0.6)
    section.bottom_margin = _inches(0.6)
    section.left_margin = _inches(0.65)
    section.right_margin = _inches(0.65)


def _add_indicators_table(document, total_feedbacks: int, negative_percent: float, consequence_counts: dict[str, int]) -> None:
    rows = (
        ("Total de Feedbacks Analisados", str(total_feedbacks)),
        ("Indice de Sentimento Negativo", f"{negative_percent:.1f}%"),
        ("Total de Correcoes", str(consequence_counts.get("Correction", 0))),
        ("Total de Melhorias", str(consequence_counts.get("Improvement", 0))),
    )
    _add_simple_table(document, ("Indicador", "Valor"), rows)


def _add_consequence_distribution_table(document, consequence_counts: dict[str, int], total_feedbacks: int) -> None:
    rows = []
    for label in CONSEQUENCE_LABELS:
        count = consequence_counts.get(label, 0)
        rows.append((label, str(count), f"{_percent(count, total_feedbacks):.1f}%"))
    _add_simple_table(document, ("Consequencia", "Casos", "Percentual"), rows)


def _add_critical_issues_table(document, records: list[FeedbackRecord]) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("ID", "Texto do Feedback", "Nivel de Sentimento", "Alvo Semantico Inferido", "Jira")
    _fill_header(table.rows[0], headers)

    for index, record in enumerate(records, start=1):
        cells = table.add_row().cells
        cells[0].text = str(record.source_id)
        cells[1].text = sanitize_feedback_text(record.text)
        cells[2].text = _sentiment_label(record.sentiment_score)
        cells[3].text = record.inferred_target or record.technical_target or "-"
        cells[4].text = _jira_status_label(record)
        if index % 2 == 0:
            _shade_row(cells, "F8FBFD")
        if cells[4].text == "Pendente de Exportacao":
            _shade_cell(cells[4], "FFF7DF")


def _add_simple_table(document, headers: tuple[str, ...], rows) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _fill_header(table.rows[0], headers)
    for index, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
        if index % 2 == 0:
            _shade_row(cells, "F8FBFD")


def _fill_header(row, headers: tuple[str, ...]) -> None:
    for cell, header in zip(row.cells, headers):
        cell.text = header
        _shade_cell(cell, "DDEBF7")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _consequence_counts(feedbacks) -> dict[str, int]:
    counts = {label: 0 for label in CONSEQUENCE_LABELS}
    for item in feedbacks.order_by().values("consequences__consequence_type").annotate(total=Count("id", distinct=True)):
        consequence = item["consequences__consequence_type"] or ""
        if consequence in counts:
            counts[consequence] = item["total"]
    for item in feedbacks.filter(consequences__isnull=True).order_by().values("consequence").annotate(total=Count("id")):
        consequence = item["consequence"] or ""
        if consequence in counts:
            counts[consequence] += item["total"]
    return counts


def _top_critical_targets(feedbacks) -> list[tuple[str, str]]:
    return [
        (item["inferred_target"], str(item["total"]))
        for item in feedbacks.filter(consequence="Correction")
        .exclude(inferred_target="")
        .order_by()
        .values("inferred_target")
        .annotate(total=Count("id"))
        .order_by("-total", "inferred_target")[:5]
    ]


def _sentiment_by_target(feedbacks) -> list[tuple[str, str]]:
    from django.db.models import Avg

    return [
        (item["inferred_target"], f"{float(item['avg_sentiment'] or 0):.3f}")
        for item in feedbacks.exclude(sentiment_score__isnull=True)
        .exclude(inferred_target="")
        .order_by()
        .values("inferred_target")
        .annotate(avg_sentiment=Avg("sentiment_score"), total=Count("id"))
        .order_by("-total", "inferred_target")[:6]
    ]


def _sentiment_label(value) -> str:
    if value is None:
        return "Nao classificado"
    score = float(value)
    if score < -0.05:
        return f"Negativo ({score:.2f})"
    if score > 0.05:
        return f"Positivo ({score:.2f})"
    return f"Neutro ({score:.2f})"


def _jira_status_label(record: FeedbackRecord) -> str:
    if record.jira_key:
        return record.jira_key
    if record.jira_status == FeedbackRecord.JiraStatus.DRY_RUN:
        return "Simulado em Dry-run"
    return "Pendente de Exportacao"


def _percent(count: int, total: int) -> float:
    return round((count / total) * 100, 1) if total else 0.0


def _inches(value: float):
    from docx.shared import Inches

    return Inches(value)


def _shade_row(cells, fill: str) -> None:
    for cell in cells:
        _shade_cell(cell, fill)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
