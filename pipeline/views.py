import csv
from io import BytesIO
from threading import Thread
from urllib.parse import urlencode

from django.conf import settings
from django.core.paginator import Paginator
from django.db import close_old_connections
from django.db.models import Avg, Count, FloatField, Q, Value
from django.db.models.functions import Coalesce
from django.http import Http404, HttpRequest, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_GET, require_POST

from .models import FeedbackRecord, PipelineEvent, ProcessingJob
from .services.ontology import FeedOnOntologyService
from .services.processor import process_job
from .tasks import process_feedback_job

ALLOWED_SENTIMENT_FILTERS = {"all", "negative", "neutral", "positive"}
ALLOWED_CONSEQUENCE_FILTERS = {"all", "Correction", "Improvement", "Prioritization"}
NEGATIVE_THRESHOLD = -0.05
POSITIVE_THRESHOLD = 0.05


@require_GET
def index(request: HttpRequest):
    jobs = ProcessingJob.objects.all()[:10]
    return render(
        request,
        "pipeline/index.html",
        {
            "jobs": jobs,
            "max_upload_size_mb": settings.MAX_UPLOAD_SIZE_MB,
            "active_nav": "upload",
        },
    )


@require_POST
def create_job(request: HttpRequest):
    upload = request.FILES.get("dataset")
    if upload is None:
        return JsonResponse({"error": "Envie um arquivo CSV."}, status=400)
    if not upload.name.lower().endswith(".csv"):
        return JsonResponse({"error": "O arquivo precisa estar no formato CSV."}, status=400)

    max_size = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if upload.size > max_size:
        return JsonResponse({"error": f"O arquivo excede {settings.MAX_UPLOAD_SIZE_MB} MB."}, status=400)

    row_limit, error = _parse_row_limit(request.POST.get("row_limit", ""))
    if error:
        return JsonResponse({"error": error}, status=400)

    job = ProcessingJob.objects.create(original_filename=upload.name, upload=upload, row_limit=row_limit)
    try:
        if settings.CELERY_TASK_ALWAYS_EAGER:
            _start_local_background_job(job.id)
            task_id = "local-thread"
        else:
            async_result = process_feedback_job.delay(job.id)
            task_id = async_result.id
    except Exception as exc:
        job.mark_failed(f"Nao foi possivel enviar o job ao Celery: {exc}")
        return JsonResponse(
            {
                "error": "Nao foi possivel iniciar o processamento. Verifique Redis/Celery ou use CELERY_TASK_ALWAYS_EAGER=true."
            },
            status=503,
        )

    job.metadata = {**job.metadata, "celery_task_id": task_id}
    job.save(update_fields=["metadata", "updated_at"])

    return JsonResponse(_job_urls(request, job), status=201)


@require_POST
def cancel_job(request: HttpRequest, job_id: int):
    job = get_object_or_404(ProcessingJob, pk=job_id)
    if job.status in {ProcessingJob.Status.COMPLETED, ProcessingJob.Status.FAILED, ProcessingJob.Status.CANCELED}:
        return JsonResponse({"status": job.status, "message": "Este job ja foi finalizado."})

    job.request_cancel()
    PipelineEvent.objects.create(job=job, level=PipelineEvent.Level.WARNING, message="Cancelamento solicitado pelo usuario.")
    return JsonResponse({"status": job.status, "message": "Cancelamento solicitado."})


@require_GET
def job_status(request: HttpRequest, job_id: int):
    job = get_object_or_404(ProcessingJob, pk=job_id)
    events = list(job.events.values("level", "message", "created_at", "metadata"))
    feedbacks = list(
        FeedbackRecord.objects.filter(job=job).values(
            "source_id",
            "text",
            "intent",
            "ai_intent",
            "sentiment_score",
            "ai_provider",
            "target_candidate",
            "technical_target",
            "inferred_target",
            "consequence",
            "jira_status",
            "jira_key",
            "processing_error",
        )[:200]
    )

    payload = {
        "id": job.id,
        "status": job.status,
        "current_phase": job.current_phase,
        "total_rows": job.total_rows,
        "processed_rows": job.processed_rows,
        "row_limit": job.row_limit,
        "cancel_requested": job.cancel_requested,
        "progress_percent": job.progress_percent,
        "jira_created": job.jira_created,
        "error_message": job.error_message,
        "events": events,
        "feedbacks": feedbacks,
    }
    payload.update(_job_urls(request, job))
    return JsonResponse(payload)


@require_GET
def dashboard(request: HttpRequest):
    jobs = list(ProcessingJob.objects.all()[:50])
    selected_job = _resolve_selected_job(request, jobs=jobs)
    sentiment_filter, consequence_filter = _parse_filters(request)
    return render(
        request,
        "pipeline/dashboard.html",
        {
            "jobs": jobs,
            "selected_job": selected_job,
            "selected_sentiment": sentiment_filter,
            "selected_consequence": consequence_filter,
            "active_nav": "dashboard",
        },
    )


@require_GET
def dashboard_data(request: HttpRequest):
    selected_job = _resolve_selected_job(request)
    if selected_job is None:
        return JsonResponse(
            {
                "job": None,
                "cards": {"total_feedbacks": 0, "negative_percent": 0, "jira_tickets": 0},
                "charts": {
                    "consequence_distribution": {"labels": [], "data": []},
                    "top_critical_features": {"labels": [], "data": []},
                    "sentiment_by_category": {"labels": [], "data": []},
                },
                "top_critical_issues": [],
                "export_urls": {"csv": "", "docx": ""},
            }
        )

    sentiment_filter, consequence_filter = _parse_filters(request)
    feedbacks = _filtered_feedbacks(selected_job, sentiment_filter, consequence_filter)
    snapshot = _build_dashboard_snapshot(request, selected_job, feedbacks, sentiment_filter, consequence_filter)
    return JsonResponse(snapshot)


@require_GET
def detailed_results(request: HttpRequest):
    jobs = list(ProcessingJob.objects.all()[:50])
    selected_job = _resolve_selected_job(request, jobs=jobs)
    sentiment_filter, consequence_filter = _parse_filters(request)

    page_obj = None
    rows = []
    if selected_job is not None:
        feedbacks = _filtered_feedbacks(selected_job, sentiment_filter, consequence_filter).order_by("-id")
        paginator = Paginator(feedbacks, 50)
        page_obj = paginator.get_page(request.GET.get("page") or 1)
        rows = list(page_obj.object_list)

    return render(
        request,
        "pipeline/detailed_results.html",
        {
            "active_nav": "results",
            "jobs": jobs,
            "selected_job": selected_job,
            "selected_sentiment": sentiment_filter,
            "selected_consequence": consequence_filter,
            "rows": rows,
            "page_obj": page_obj,
            "base_query": _query_without_page(request),
        },
    )


@require_GET
def export_csv(request: HttpRequest):
    job = _require_selected_job(request)
    sentiment_filter, consequence_filter = _parse_filters(request)
    feedbacks = _filtered_feedbacks(job, sentiment_filter, consequence_filter).order_by("id")
    ontology = FeedOnOntologyService()

    response = HttpResponse(content_type="text/csv; charset=utf-8")
    filename = f"feed-on-job-{job.id}-resultados.csv"
    response["Content-Disposition"] = f'attachment; filename="{filename}"'

    writer = csv.writer(response)
    writer.writerow(
        [
            "ID",
            "Texto Original",
            "Sentimento",
            "Alvo IA",
            "Alvo Inferido",
            "Consequencia",
            "Link do Jira",
        ]
    )

    for record in feedbacks.iterator(chunk_size=500):
        writer.writerow(
            [
                record.source_id,
                record.text,
                _format_sentiment(record.sentiment_score),
                record.target_candidate,
                _semantic_inferred_target(ontology, record),
                record.consequence,
                _jira_issue_url(record.jira_key),
            ]
        )

    return response


@require_GET
def export_docx(request: HttpRequest):
    job = _require_selected_job(request)
    sentiment_filter, consequence_filter = _parse_filters(request)
    feedbacks = _filtered_feedbacks(job, sentiment_filter, consequence_filter).order_by("id")
    snapshot = _build_dashboard_snapshot(request, job, feedbacks, sentiment_filter, consequence_filter)
    ontology = FeedOnOntologyService()

    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        return JsonResponse({"error": f"Biblioteca python-docx indisponivel: {exc}"}, status=503)

    document = Document()
    run_date = timezone.localtime(job.finished_at or timezone.now()).strftime("%d/%m/%Y %H:%M")

    document.add_heading("Relatorio Executivo FEED-ON", level=1)
    document.add_paragraph(f"Data do processamento: {run_date}")
    document.add_paragraph(f"Lote analisado: Job #{job.id} - {job.original_filename}")

    cards = snapshot["cards"]
    consequence_chart = snapshot["charts"]["consequence_distribution"]
    critical_chart = snapshot["charts"]["top_critical_features"]
    sentiment_chart = snapshot["charts"]["sentiment_by_category"]

    document.add_heading("Resumo Executivo", level=2)
    document.add_paragraph(
        (
            f"Foram avaliados {cards['total_feedbacks']} feedbacks na visao atual. "
            f"{cards['negative_percent']}% apresentam sentimento negativo e "
            f"{cards['jira_tickets']} tickets Jira foram gerados."
        )
    )
    document.add_paragraph(
        "Distribuicao de consequencias: "
        + _chart_textual_summary(consequence_chart["labels"], consequence_chart["data"], "%")
    )
    document.add_paragraph(
        "Top 5 features criticas (concentracao de Correction): "
        + _chart_textual_summary(critical_chart["labels"], critical_chart["data"], "casos")
    )
    document.add_paragraph(
        "Media de sentimento por categoria: "
        + _chart_textual_summary(sentiment_chart["labels"], sentiment_chart["data"], "media")
    )

    critical_rows = list(
        feedbacks.filter(consequence="Correction")
        .annotate(sentiment_sort=Coalesce("sentiment_score", Value(1.0), output_field=FloatField()))
        .order_by("sentiment_sort", "id")[:10]
    )
    document.add_heading("Top 10 Critical Issues", level=2)
    if not critical_rows:
        document.add_paragraph("Nao ha issues criticos para os filtros selecionados.")
    else:
        for record in critical_rows:
            document.add_paragraph(
                (
                    f"[{record.source_id}] {record.text} | "
                    f"Sentimento: {_format_sentiment(record.sentiment_score)} | "
                    f"Alvo Inferido: {_semantic_inferred_target(ontology, record)} | "
                    f"Jira: {record.jira_key or 'nao gerado'}"
                ),
                style="List Number",
            )

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    filename = f"feed-on-job-{job.id}-relatorio-executivo.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _parse_row_limit(raw_value: str) -> tuple[int | None, str]:
    raw_value = (raw_value or "").strip()
    if not raw_value:
        return None, ""
    try:
        value = int(raw_value)
    except ValueError:
        return None, "O limite precisa ser um numero inteiro."
    if value <= 0:
        return None, "O limite precisa ser maior que zero."
    return value, ""


def _job_urls(request: HttpRequest, job: ProcessingJob) -> dict:
    return {
        "job_id": job.id,
        "status_url": request.build_absolute_uri(reverse("pipeline:job_status", args=[job.id])),
        "cancel_url": request.build_absolute_uri(reverse("pipeline:cancel_job", args=[job.id])),
    }


def _start_local_background_job(job_id: int) -> None:
    def runner() -> None:
        close_old_connections()
        try:
            process_job(job_id)
        finally:
            close_old_connections()

    Thread(target=runner, daemon=True).start()


def _resolve_selected_job(request: HttpRequest, jobs: list[ProcessingJob] | None = None) -> ProcessingJob | None:
    selected_job_id = request.GET.get("job")
    if selected_job_id:
        try:
            return ProcessingJob.objects.get(pk=int(selected_job_id))
        except (ValueError, ProcessingJob.DoesNotExist):
            pass

    if jobs is not None:
        return jobs[0] if jobs else None

    return ProcessingJob.objects.first()


def _require_selected_job(request: HttpRequest) -> ProcessingJob:
    selected_job = _resolve_selected_job(request)
    if selected_job is None:
        raise Http404("Nenhum job disponivel para exportacao.")
    return selected_job


def _parse_filters(request: HttpRequest) -> tuple[str, str]:
    sentiment_filter = (request.GET.get("sentiment") or "all").strip().lower()
    if sentiment_filter not in ALLOWED_SENTIMENT_FILTERS:
        sentiment_filter = "all"

    consequence_filter = (request.GET.get("consequence") or "all").strip()
    if consequence_filter not in ALLOWED_CONSEQUENCE_FILTERS:
        consequence_filter = "all"

    return sentiment_filter, consequence_filter


def _filtered_feedbacks(job: ProcessingJob, sentiment_filter: str, consequence_filter: str):
    feedbacks = FeedbackRecord.objects.filter(job=job)

    if sentiment_filter == "negative":
        feedbacks = feedbacks.filter(sentiment_score__lt=NEGATIVE_THRESHOLD)
    elif sentiment_filter == "neutral":
        feedbacks = feedbacks.filter(sentiment_score__gte=NEGATIVE_THRESHOLD, sentiment_score__lte=POSITIVE_THRESHOLD)
    elif sentiment_filter == "positive":
        feedbacks = feedbacks.filter(sentiment_score__gt=POSITIVE_THRESHOLD)

    if consequence_filter != "all":
        feedbacks = feedbacks.filter(consequence=consequence_filter)

    return feedbacks


def _build_dashboard_snapshot(
    request: HttpRequest,
    job: ProcessingJob,
    feedbacks,
    sentiment_filter: str,
    consequence_filter: str,
) -> dict:
    total_feedbacks = feedbacks.count()
    negative_count = feedbacks.filter(sentiment_score__lt=NEGATIVE_THRESHOLD).count()
    negative_percent = round((negative_count / total_feedbacks) * 100, 1) if total_feedbacks else 0
    jira_tickets = feedbacks.filter(
        Q(jira_status=FeedbackRecord.JiraStatus.CREATED) | Q(jira_status=FeedbackRecord.JiraStatus.DRY_RUN)
    ).count()

    consequence_counts = {label: 0 for label in ("Correction", "Improvement", "Prioritization")}
    for item in feedbacks.values("consequence").annotate(total=Count("id")):
        consequence = item["consequence"] or ""
        if consequence in consequence_counts:
            consequence_counts[consequence] = item["total"]

    if total_feedbacks:
        consequence_distribution = [round((value / total_feedbacks) * 100, 1) for value in consequence_counts.values()]
    else:
        consequence_distribution = [0, 0, 0]

    top_critical = list(
        feedbacks.filter(consequence="Correction")
        .exclude(inferred_target="")
        .values("inferred_target")
        .annotate(total=Count("id"))
        .order_by("-total", "inferred_target")[:5]
    )

    sentiment_by_category = list(
        feedbacks.exclude(sentiment_score__isnull=True)
        .exclude(inferred_target="")
        .values("inferred_target")
        .annotate(avg_sentiment=Avg("sentiment_score"), total=Count("id"))
        .order_by("-total", "inferred_target")[:6]
    )

    critical_issues = list(
        feedbacks.filter(consequence="Correction")
        .annotate(sentiment_sort=Coalesce("sentiment_score", Value(1.0), output_field=FloatField()))
        .order_by("sentiment_sort", "id")
        .values("source_id", "text", "sentiment_score", "inferred_target", "jira_key")[:10]
    )

    filter_query = urlencode({"job": job.id, "sentiment": sentiment_filter, "consequence": consequence_filter})
    csv_url = f"{reverse('pipeline:export_csv')}?{filter_query}"
    docx_url = f"{reverse('pipeline:export_docx')}?{filter_query}"

    return {
        "job": {
            "id": job.id,
            "status": job.status,
            "filename": job.original_filename,
            "processed_rows": job.processed_rows,
            "total_rows": job.total_rows,
            "finished_at": timezone.localtime(job.finished_at).isoformat() if job.finished_at else None,
        },
        "filters": {"sentiment": sentiment_filter, "consequence": consequence_filter},
        "cards": {
            "total_feedbacks": total_feedbacks,
            "negative_percent": negative_percent,
            "jira_tickets": jira_tickets,
        },
        "charts": {
            "consequence_distribution": {
                "labels": list(consequence_counts.keys()),
                "data": consequence_distribution,
            },
            "top_critical_features": {
                "labels": [item["inferred_target"] for item in top_critical],
                "data": [item["total"] for item in top_critical],
            },
            "sentiment_by_category": {
                "labels": [item["inferred_target"] for item in sentiment_by_category],
                "data": [round(item["avg_sentiment"] or 0, 3) for item in sentiment_by_category],
            },
        },
        "top_critical_issues": [
            {
                "source_id": item["source_id"],
                "text": item["text"],
                "sentiment_score": _format_sentiment(item["sentiment_score"]),
                "inferred_target": item["inferred_target"] or "-",
                "jira_key": item["jira_key"] or "-",
                "jira_url": _jira_issue_url(item["jira_key"]),
            }
            for item in critical_issues
        ],
        "export_urls": {
            "csv": request.build_absolute_uri(csv_url),
            "docx": request.build_absolute_uri(docx_url),
        },
    }


def _semantic_inferred_target(ontology: FeedOnOntologyService, record: FeedbackRecord) -> str:
    source_id = f"{record.source_id}_{record.id}"
    technical_target = record.technical_target or record.inferred_target or "Feature.General"

    try:
        result = ontology.interpret(
            source_id=source_id,
            text=record.text,
            intent=record.intent,
            technical_target=technical_target,
        )
        semantic_value = ontology.inferred_target_for(source_id)
        return semantic_value or result.inferred_target or record.inferred_target or technical_target
    except Exception:
        return record.inferred_target or technical_target


def _jira_issue_url(jira_key: str) -> str:
    if not jira_key:
        return ""
    base = (settings.JIRA_URL or "").rstrip("/")
    if not base:
        return jira_key
    return f"{base}/browse/{jira_key}"


def _format_sentiment(value) -> str:
    if value is None:
        return "-"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return "-"


def _chart_textual_summary(labels: list[str], values: list[float], suffix: str) -> str:
    if not labels:
        return "sem dados para o filtro aplicado."

    parts = []
    for label, value in zip(labels, values):
        if isinstance(value, float):
            formatted = f"{value:.1f}"
        else:
            formatted = str(value)
        parts.append(f"{label}: {formatted} {suffix}".strip())
    return "; ".join(parts)


def _query_without_page(request: HttpRequest) -> str:
    query = request.GET.copy()
    query.pop("page", None)
    return query.urlencode()
